import re
import sys
import requests
from docx import Document
# ─────────────────────────────────────────────
#  CONFIGURACIÓN
# ─────────────────────────────────────────────
ACCESS_TOKEN = "AQUÍ VA EL TOKEN" # CAMBIAR SOLO CON EL TOKEN
SURVEY_TITLE = "NOMBRE DE LA ENCUESTA"# CAMBIAR A GUSTO
DOCX_PATH    = "AQUÍ LA RUTA DEL DOCUMENTO DOCX" #C CAMBIAR A GUSTO
PAGE_TITLE   = "Sección 1" # CAMBIAR A GUSTO
# Se ajusta la escala a los valores numéricos que pide el cuestionario
LIKERT_SCALE = ["0", "1", "2", "3", "4"] 
# ─────────────────────────────────────────────

# CLAVE
BASE_URL = "https://api.surveymonkey.com/v3"
HEADERS  = {
    "Authorization": "Bearer " + ACCESS_TOKEN,
    "Content-Type":  "application/json",
}

# ──────────────────────────────────────────────
#  1. LECTURA Y PARSEO DEL DOCX (ADAPTADO)
# ──────────────────────────────────────────────

def leer_docx(path):
    doc = Document(path)
    preguntas = []
    pregunta_actual = None
    opciones_actuales = []
    
    # Expresión para detectar las líneas en blanco para rellenar
    re_blancos = re.compile(r"_{3,}")

    def guardar_pregunta():
        # Filtramos instrucciones largas que no son opciones
        opciones_limpias = [op for op in opciones_actuales if len(op) < 150]
        if pregunta_actual and pregunta_actual.strip():
            tipo = detectar_tipo(pregunta_actual, opciones_limpias)
            preguntas.append({
                "texto":    pregunta_actual.strip(),
                "tipo":     tipo,
                "opciones": list(opciones_limpias),
            })

    # A. Extracción de preguntas estándar desde los párrafos
    in_intro = True
    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        # Saltar la introducción hasta llegar a la primera sección
        if "SECCIÓN I" in texto:
            in_intro = False
            continue
        # Ignorar subtítulos de sección o textos introductorios adicionales
        if in_intro or texto.startswith("SECCIÓN") or texto.startswith("INTRODUCCIÓN"):
            continue

        # Heurística: Es pregunta si termina en ?, :, tiene "____" o es "Sexo"
        es_pregunta = (
            texto.endswith("?") or 
            texto.endswith(":") or 
            bool(re_blancos.search(texto)) or 
            texto.lower() == "sexo"
        )

        if es_pregunta:
            guardar_pregunta()
            # Limpiamos las rayas de la pregunta para que quede limpia en SurveyMonkey
            pregunta_actual = re_blancos.sub("", texto).strip()
            opciones_actuales = []
        elif pregunta_actual:
            # Si hay una pregunta activa y el texto no es una pregunta nueva, es una opción
            opciones_actuales.append(texto)

    guardar_pregunta()

    # B. Extracción de matrices (Likert) desde las tablas
    for table in doc.tables:
        for row in table.rows:
            celda_texto = row.cells[0].text.strip()
            # Si el texto es suficientemente largo y no es el encabezado de la tabla
            if len(celda_texto) > 20 and not celda_texto.startswith("Objetivos"):
                preguntas.append({
                    "texto": celda_texto,
                    "tipo": "likert",
                    "opciones": [] # Las opciones se manejan con LIKERT_SCALE en el payload
                })

    return preguntas

def detectar_tipo(texto, opciones):
    # Ya no dependemos de palabras clave para Likert, pues vienen de las tablas
    if opciones:
        return "multiple"
    return "abierta"

# ──────────────────────────────────────────────
#  2. TRANSFORMACIÓN AL FORMATO DE SURVEYMONKEY
# ──────────────────────────────────────────────

def construir_payload_pregunta(pregunta, posicion):
    tipo = pregunta["tipo"]

    if tipo == "multiple":
        return {
            "headings": [{"heading": pregunta["texto"]}],
            "family":   "single_choice",
            "subtype":  "vertical",
            "position": posicion,
            "answers": {
                "choices": [{"text": op} for op in pregunta["opciones"]]
            },
        }

    elif tipo == "likert":
        return {
            "headings": [{"heading": pregunta["texto"]}],
            "family":   "matrix",
            "subtype":  "rating",
            "position": posicion,
            "answers": {
                "rows":    [{"text": ""}],
                "choices": [{"text": et} for et in LIKERT_SCALE],
            },
        }

    else:
        return {
            "headings": [{"heading": pregunta["texto"]}],
            "family":   "open_ended",
            "subtype":  "essay",
            "position": posicion,
        }


# ──────────────────────────────────────────────
#  3. LLAMADAS A LA API DE SURVEYMONKEY
# ──────────────────────────────────────────────

def crear_encuesta(titulo):
    resp = requests.post(
        BASE_URL + "/surveys",
        headers=HEADERS,
        json={"title": titulo},
    )
    resp.raise_for_status()
    survey_id = resp.json()["id"]
    print("Encuesta creada: ID=" + str(survey_id))
    return survey_id


def crear_pagina(survey_id, titulo):
    resp = requests.post(
        BASE_URL + "/surveys/" + str(survey_id) + "/pages",
        headers=HEADERS,
        json={"title": titulo},
    )
    resp.raise_for_status()
    page_id = resp.json()["id"]
    print("Página creada: ID=" + str(page_id))
    return page_id


def agregar_pregunta(survey_id, page_id, payload):
    resp = requests.post(
        BASE_URL + "/surveys/" + str(survey_id) + "/pages/" + str(page_id) + "/questions",
        headers=HEADERS,
        json=payload,
    )
    if resp.status_code not in (200, 201):
        pos = payload["position"]
        print("  Error en pregunta pos=" + str(pos) + ": " + resp.text)
    else:
        familia = payload.get("family", "?")
        heading = payload["headings"][0]["heading"][:60]
        print("  [" + familia + "] " + heading)


# ──────────────────────────────────────────────
#  4. FLUJO PRINCIPAL
# ──────────────────────────────────────────────

def main():
    print("Leyendo cuestionario: " + DOCX_PATH)
    preguntas = leer_docx(DOCX_PATH)

    if not preguntas:
        print("No se encontraron preguntas. Las preguntas deben iniciar con número: '1. Pregunta'")
        sys.exit(1)

    total      = len(preguntas)
    n_multiple = sum(1 for p in preguntas if p["tipo"] == "multiple")
    n_likert   = sum(1 for p in preguntas if p["tipo"] == "likert")
    n_abierta  = sum(1 for p in preguntas if p["tipo"] == "abierta")
    print("Encontradas " + str(total) + " preguntas:")
    print("  Opcion multiple : " + str(n_multiple))
    print("  Escala Likert   : " + str(n_likert))
    print("  Abiertas        : " + str(n_abierta))

    print("Subiendo a SurveyMonkey...")
    survey_id = crear_encuesta(SURVEY_TITLE)
    page_id   = crear_pagina(survey_id, PAGE_TITLE)

    for i, pregunta in enumerate(preguntas, start=1):
        payload = construir_payload_pregunta(pregunta, posicion=i)
        agregar_pregunta(survey_id, page_id, payload)

    print("Listo! Encuesta disponible en:")
    print("https://www.surveymonkey.com/create/?sm=" + str(survey_id))


if __name__ == "__main__":
    main()